from __future__ import annotations

import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from .contracts import (
    KNOWLEDGE_CATEGORIES,
    SUPPORTED_DOCUMENT_EXTENSIONS,
    DocumentParseError,
    KnowledgeParseSummary,
    ParsedDocumentMetadata,
    ParsedKnowledgeDocument,
)


PARSER_VERSION = "1.0.0"


class KnowledgeDocumentParserError(Exception):
    pass

#不支持的文件类型错误
class UnsupportedDocumentTypeError(KnowledgeDocumentParserError):
    pass

#空文件错误
class EmptyDocumentError(KnowledgeDocumentParserError):
    pass


class KnowledgeDocumentParser:
    def __init__(self, knowledge_root: str | Path) -> None:
        self.knowledge_root = Path(knowledge_root).resolve() #把这个“路径对象”变成“绝对路径”

    def parse_file(self, file_path: str | Path) -> ParsedKnowledgeDocument:
        path = Path(file_path).resolve()
        category = self._category_for_path(path)
        extension = path.suffix.lower() #获取文件后缀的小写

        if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
            raise UnsupportedDocumentTypeError(f"Unsupported document type: {extension}")
        if not path.is_file():
            raise KnowledgeDocumentParserError(f"Knowledge document does not exist: {path}")

        raw_bytes = path.read_bytes()
        content_hash = hashlib.sha256(raw_bytes).hexdigest()
        content = self._parse_content(path, raw_bytes)
        normalized_content = self._normalize_content(content)
        if not normalized_content:
            raise EmptyDocumentError("Parsed document content is empty")

        stat = path.stat()
        manifest = self._load_manifest(path)
        metadata = ParsedDocumentMetadata(
            file_path=str(path),
            file_name=path.name,
            file_type=extension.lstrip("."),
            file_size_bytes=stat.st_size,
            content_hash=content_hash,
            modified_at=datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc),
            parser="KnowledgeDocumentParser",
            parser_version=PARSER_VERSION,
            title=self._optional_text(manifest.get("title")) or path.stem,
            publisher=self._optional_text(manifest.get("publisher")),
            source_url=self._optional_text(manifest.get("source_url")),
            document_version=self._optional_text(manifest.get("version")),
            reviewed_at=self._optional_datetime(manifest.get("reviewed_at")),
            parent_category=self._optional_text(manifest.get("parent_category")),
            child_categories=self._string_list(manifest.get("child_categories")),
            applicable_population=self._string_list(manifest.get("applicable_population")),
            evidence_level=self._optional_text(manifest.get("evidence_level")),
            medical_boundary=self._optional_text(manifest.get("medical_boundary")),
            source_type=self._optional_text(manifest.get("source_type")),
        )

        return ParsedKnowledgeDocument(
            document_id=self._document_id(category, path, self.knowledge_root),
            source=self._optional_text(manifest.get("title")) or path.stem,
            category=category,
            content=normalized_content,
            metadata=metadata,
        )

    def parse_all(self) -> KnowledgeParseSummary:
        documents: list[ParsedKnowledgeDocument] = []
        errors: list[DocumentParseError] = []
        skipped_files: list[str] = []
        total_files = 0

        for category in KNOWLEDGE_CATEGORIES:
            category_dir = self.knowledge_root / category
            if not category_dir.exists():
                continue

            for path in sorted(category_dir.rglob("*")):
                if path.is_dir():
                    continue
                total_files += 1
                if path.suffix.lower() not in SUPPORTED_DOCUMENT_EXTENSIONS:
                    skipped_files.append(str(path))
                    continue
                try:
                    documents.append(self.parse_file(path))
                except Exception as exc:  # noqa: BLE001 - converted to structured parse error.
                    errors.append(
                        DocumentParseError(
                            file_path=str(path),
                            category=category,
                            error_code=exc.__class__.__name__,
                            message=str(exc),
                        )
                    )

        return KnowledgeParseSummary(
            root_path=str(self.knowledge_root),
            total_files=total_files,
            parsed_count=len(documents),
            failed_count=len(errors),
            skipped_count=len(skipped_files),
            documents=documents,
            errors=errors,
            skipped_files=skipped_files,
        )

    def _category_for_path(self, path: Path) -> str:
        try:
            relative = path.relative_to(self.knowledge_root)
        except ValueError as exc:
            raise KnowledgeDocumentParserError(
                f"Knowledge document must be under {self.knowledge_root}"
            ) from exc

        if not relative.parts:
            raise KnowledgeDocumentParserError("Knowledge document has no category directory")

        category = relative.parts[0]
        if category not in KNOWLEDGE_CATEGORIES:
            raise KnowledgeDocumentParserError(f"Unknown knowledge category: {category}")
        return category

    def _parse_content(self, path: Path, raw_bytes: bytes) -> str:
        extension = path.suffix.lower()
        if extension in {".md", ".markdown", ".txt"}:
            return self._decode_text(raw_bytes)
        if extension == ".pdf":
            return self._parse_pdf(path)
        if extension == ".docx":
            return self._parse_docx(path)
        raise UnsupportedDocumentTypeError(f"Unsupported document type: {extension}")

    @staticmethod
    def _load_manifest(path: Path) -> dict[str, object]:
        candidates = (path.with_suffix(path.suffix + ".meta.json"), path.with_suffix(".meta.json"))
        for candidate in candidates:
            if not candidate.is_file():
                continue
            try:
                value = json.loads(candidate.read_text(encoding="utf-8"))
            except (OSError, UnicodeDecodeError, json.JSONDecodeError) as exc:
                raise KnowledgeDocumentParserError(f"Invalid metadata manifest {candidate}: {exc}") from exc
            if not isinstance(value, dict):
                raise KnowledgeDocumentParserError(f"Metadata manifest must be an object: {candidate}")
            return value
        return {}

    @staticmethod
    def _optional_text(value: object) -> str | None:
        text = str(value).strip() if value is not None else ""
        return text or None

    @staticmethod
    def _optional_datetime(value: object) -> datetime | None:
        if value is None or not str(value).strip():
            return None
        try:
            parsed = datetime.fromisoformat(str(value).replace("Z", "+00:00"))
        except ValueError as exc:
            raise KnowledgeDocumentParserError(f"reviewed_at must be ISO-8601: {value}") from exc
        return parsed if parsed.tzinfo else parsed.replace(tzinfo=timezone.utc)

    @staticmethod
    def _string_list(value: object) -> list[str]:
        if not isinstance(value, list):
            return []
        return [str(item).strip() for item in value if str(item).strip()]

    @staticmethod
    def _decode_text(raw_bytes: bytes) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return raw_bytes.decode(encoding)
            except UnicodeDecodeError:
                continue
        return raw_bytes.decode("utf-8", errors="replace")

    @staticmethod
    def _parse_pdf(path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise KnowledgeDocumentParserError(
                "Missing dependency pypdf. Install requirements before parsing PDF files."
            ) from exc

        reader = PdfReader(str(path))
        pages: list[str] = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n\n".join(pages)

    @staticmethod
    def _parse_docx(path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise KnowledgeDocumentParserError(
                "Missing dependency python-docx. Install requirements before parsing DOCX files."
            ) from exc

        document = Document(str(path))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        return "\n\n".join(blocks)

    @staticmethod
    def _normalize_content(content: str) -> str:
        content = content.replace("\r\n", "\n").replace("\r", "\n")
        content = re.sub(r"[ \t]+", " ", content)
        content = re.sub(r"\n{3,}", "\n\n", content)
        return content.strip()

    @staticmethod
    def _document_id(category: str, path: Path, knowledge_root: Path) -> str:
        safe_stem = re.sub(r"[^a-zA-Z0-9_-]+", "-", path.stem).strip("-").lower()  
        #path.stem：去掉文件的扩展名，txt,pdf,docx等
        if not safe_stem:
            safe_stem = "document"

        #获取文件的相对路径，path到knowledge_root的相对路径    
        relative_key = path.relative_to(knowledge_root).as_posix().lower()
        path_hash = hashlib.sha256(relative_key.encode("utf-8")).hexdigest()[:12]
        return f"{category}-{safe_stem}-{path_hash}"
